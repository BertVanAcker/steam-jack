###################################################################################
#	Author:			Bert Van Acker (bva.bmkr@gmail.com)
#	Version:		0.1.0
#	Lisence:		LGPL-3.0 (GNU Lesser General Public License version 3)
#
#	Description:	Documentor class handling document and template generation
###################################################################################

#imports
import nbformat as nbf
from os import mkdir
from os.path import exists, dirname, join
import jinja2
from docx import Document
from docx.shared import Mm
import python_markdown_maker


class Documentor():
    """
        Documentor: Class representing the generic documentation handler

         :param bool DEBUG: setting the verbose
         :param object Logger: Logger object for uniform logging
    """
    def __init__(self,DEBUG=True,LOGGER=None):

        #verbose and logging
        self.DEBUG = DEBUG
        self.LOGGER = LOGGER


    def generateNotebook(self,NotebookObject,output):
        """
            Function to generate a jupyter notebook

             :param object NotebookObject: Jupyter notebook class
             :param string output: Path to the output location
        """

        self.notebook = nbf.v4.new_notebook()

        content = []
        for cell in NotebookObject.Content:
            if cell.Format=='intro':
                title = '# '+cell.Title+'\n'  #heading 1 title
            elif cell.Format=='subsection':
                title = '## ' + cell.Title + '\n'  # heading 2 title
            elif cell.Format=='subsubsection':
                title = '### ' + cell.Title + '\n'  # heading 3 title
            text = cell.Text
            if cell.Type=='markdown':
                section = title+text
                content.append(nbf.v4.new_markdown_cell(section))
            if cell.Type == 'code':
                section = title+cell.Text
                codeSection = cell.Code
                content.append(nbf.v4.new_markdown_cell(section))
                content.append(nbf.v4.new_code_cell(codeSection))

        #add all cells to the notebook
        self.notebook['cells'] = content

        #Generate the jupyter notebook and store in output
        fname = output+NotebookObject.Name+'.ipynb'
        with open(fname, 'w') as f:
            nbf.write(self.notebook, f)


    def generateSTEAM_notebook(self,SteamObject,output):
        """
            Function to generate a jupyter notebook STEAM lession

             :param object SteamObject: SteamObject class
             :param string output: Path to the output location
        """

        self.notebook = nbf.v4.new_notebook()
        content = []

        #----------add the intro------------
        title = '# ' + SteamObject.Introduction.Title + '\n'  # heading 1 title
        text = SteamObject.Introduction.Text
        content.append(nbf.v4.new_markdown_cell(title + text))
        # ----------add the supported devices------------
        title = '## ' + SteamObject.SupportedDevices.Title + '\n'  # heading 2 title
        text = SteamObject.SupportedDevices.Text
        content.append(nbf.v4.new_markdown_cell(title + text))
        # ----------add the related modules------------
        title = '## ' + SteamObject.RelatedModules.Title + '\n'  # heading 2 title
        text = SteamObject.RelatedModules.Text
        content.append(nbf.v4.new_markdown_cell(title + text))
        # ----------add other content as formatted------------
        for cell in SteamObject.Content:
            if cell.Format=='intro':
                title = '# '+cell.Title+'\n'  #heading 1 title
            elif cell.Format=='subsection':
                title = '## ' + cell.Title + '\n'  # heading 2 title
            elif cell.Format=='subsubsection':
                title = '### ' + cell.Title + '\n'  # heading 3 title
            text = cell.Text
            if cell.Type=='markdown':
                section = title+text
                content.append(nbf.v4.new_markdown_cell(section))
            if cell.Type == 'code':
                section = title+cell.Text
                codeSection = cell.Code
                content.append(nbf.v4.new_markdown_cell(section))
                content.append(nbf.v4.new_code_cell(codeSection))

        #add all cells to the notebook
        self.notebook['cells'] = content

        #Generate the jupyter notebook and store in output
        fname = output+SteamObject.Name+'.ipynb'
        with open(fname, 'w') as f:
            nbf.write(self.notebook, f)

    def generateFirmwareTemplate_python(self,FirmwareObject,output):

        # Create the output folder
        srcgen_folder = join(output, 'Generated')
        if not exists(srcgen_folder):
            mkdir(srcgen_folder)

        # Initialize the template engine.
        jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader('/Users/bertvanacker/Documents/BertVanAcker/03_RemoteRepositories/steam-jack/steam_jack/'),
            trim_blocks=True,
            lstrip_blocks=True)

        # Load the Java template
        template = jinja_env.get_template('Documentor/templates/Firmware_Python_DeviceFunctions.template')

        # Generate  device-specific part of python firmware prototype

        with open((join(srcgen_folder, 'DeviceSpecific.py')), 'w') as f:
            f.write(template.render(object=FirmwareObject))

    def generateDocx(self,docList,output):

        #start a new docx document
        document = Document()

        for doc in docList:
            for docElement in doc.Content:
                if 'Header' in str(type(docElement)):
                    document.add_heading(docElement.Text, docElement.Level)
                if 'Paragraph' in str(type(docElement)):
                    document.add_paragraph(docElement.Text)
                if 'Image' in str(type(docElement)):
                    document.add_picture(docElement.ImagePath, width=Mm(int(docElement.Width)),height=Mm(int(docElement.Height)))
                if 'ListSection' in str(type(docElement)):
                    if docElement.Bullet == "-":
                        for listItem in docElement.ListElements:
                            document.add_paragraph(listItem, style='List Bullet')
                    elif docElement.Bullet=="1.":
                        document.add_paragraph(listItem, style='List Number')
                if 'CodeSection' in str(type(docElement)):
                    document.add_paragraph(docElement.Code)

        document.save(output)

    def generateMarkdown(self,docList,output):

        #start a new md document
        document = python_markdown_maker.Document()
        content = []
        for doc in docList:
            for docElement in doc.Content:
                if 'Header' in str(type(docElement)):
                    content.append(python_markdown_maker.headers(docElement.Text, level=docElement.Level))
                    content.append('\n')
                if 'Paragraph' in str(type(docElement)):
                    content.append(docElement.Text)
                    content.append('\n')
                if 'Image' in str(type(docElement)):
                    # RESOLVE github asset location
                    githubPath = "https://github.com/BertVanAcker/steam-jack/blob/main/Resources"
                    path = docElement.ImagePath.split("Resources")
                    remotePath = githubPath+path[1]

                    #content.append('\n<img src="'+remotePath+'?raw=True" width="'+docElement.Width+'" height="'+docElement.Height+'" />\n')     #TODO: no width and height setting now, check if needed!
                    content.append('\n<img src="' + remotePath + '?raw=True"/>\n')  # TODO: no width and height setting now, check if needed!
                    content.append('\n')
                if 'ListSection' in str(type(docElement)):
                    if docElement.Bullet == "-":
                        content.append(python_markdown_maker.lists(docElement.ListElements))
                        content.append('\n')
                    elif docElement.Bullet=="1.":
                        listContent = ['order']
                        for listItem in docElement.ListElements:
                            content.append(listContent.append(listItem))
                        content.append(python_markdown_maker.lists(listContent))
                        content.append('\n')
                if 'CodeSection' in str(type(docElement)):
                    content.append(python_markdown_maker.code_block(docElement.Code, lang=docElement.Formalism))
                    content.append('\n')

        #clean content
        content_cleaned = []
        for val in content:
            if val != None:
                content_cleaned.append(val)
        document.write(content_cleaned)

        with open(output, 'w') as file:
            document.render.save_as_md(file)

